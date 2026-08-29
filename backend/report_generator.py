import io
import csv
from datetime import datetime, timezone
from typing import Dict, Any, List

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_pdf_report(scan_data: Dict[str, Any]) -> io.BytesIO:
    """Generates an official Government of India Legal Metrology Inspection Compliance PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    # Custom Styles
    primary_color = colors.HexColor("#001255")
    saffron_color = colors.HexColor("#E88A1E")
    success_color = colors.HexColor("#28A745")
    danger_color = colors.HexColor("#DC3545")
    warning_color = colors.HexColor("#FFC107")
    dark_slate = colors.HexColor("#212529")
    light_bg = colors.HexColor("#F8F9FA")
    
    title_style = ParagraphStyle(
        "OfficialTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=18,
        textColor=primary_color,
        alignment=1
    )
    
    subtitle_style = ParagraphStyle(
        "OfficialSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#555555"),
        alignment=1
    )
    
    section_heading = ParagraphStyle(
        "SectionHead",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=primary_color,
        spaceAfter=4
    )
    
    cell_bold = ParagraphStyle("CellBold", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=dark_slate)
    cell_normal = ParagraphStyle("CellNormal", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=dark_slate)
    cell_danger = ParagraphStyle("CellDanger", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=danger_color)
    cell_success = ParagraphStyle("CellSuccess", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=success_color)

    # 1. Government Header Banner
    story.append(Paragraph("GOVERNMENT OF INDIA", title_style))
    story.append(Paragraph("MINISTRY OF CONSUMER AFFAIRS, FOOD & PUBLIC DISTRIBUTION", subtitle_style))
    story.append(Paragraph("DEPARTMENT OF CONSUMER AFFAIRS • LEGAL METROLOGY ENFORCEMENT WING", subtitle_style))
    story.append(Paragraph("STATUTORY PACKAGED COMMODITIES INSPECTION REPORT", ParagraphStyle("RepTitle", parent=title_style, fontSize=12, leading=15, textColor=saffron_color, spaceBefore=4)))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=2, color=primary_color, spaceAfter=8))
    
    # 2. Case Summary & Metadata Table
    scan_id = scan_data.get("id", "SCAN-0000")
    brand_name = scan_data.get("brand_name", "N/A")
    commodity_name = scan_data.get("commodity_name", "N/A")
    barcode = scan_data.get("barcode_gtin", "N/A")
    inspector_name = scan_data.get("inspector_name", "Field Officer INS-782")
    created_at = scan_data.get("created_at", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"))
    status_val = scan_data.get("compliance_status", "Non-Compliant")
    score_val = scan_data.get("compliance_score", 0)
    
    meta_data = [
        [Paragraph("<b>Inspection Case ID:</b>", cell_bold), Paragraph(str(scan_id), cell_normal), Paragraph("<b>Inspection Date:</b>", cell_bold), Paragraph(str(created_at)[:19], cell_normal)],
        [Paragraph("<b>Brand / Product:</b>", cell_bold), Paragraph(str(brand_name), cell_normal), Paragraph("<b>Commodity Type:</b>", cell_bold), Paragraph(str(commodity_name), cell_normal)],
        [Paragraph("<b>Barcode / GTIN:</b>", cell_bold), Paragraph(str(barcode), cell_normal), Paragraph("<b>Enforcement Officer:</b>", cell_bold), Paragraph(str(inspector_name), cell_normal)],
        [Paragraph("<b>Compliance Status:</b>", cell_bold), Paragraph(f"<font color='{'#28A745' if status_val == 'Compliant' else '#DC3545'}'><b>{status_val}</b></font>", cell_bold), Paragraph("<b>Compliance Score:</b>", cell_bold), Paragraph(f"<b>{score_val} / 100</b>", cell_bold)],
    ]
    meta_table = Table(meta_data, colWidths=[110, 150, 110, 150])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), light_bg),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#DDE1FF")),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 10))
    
    # 3. Rule 6 Mandatory Declarations Table
    story.append(Paragraph("1. Rule 6 Statutory Mandatory Declarations Verification", section_heading))
    decl = scan_data.get("declarations", {})
    
    decl_rows = [
        [Paragraph("<b>Statutory Field (Rule 6)</b>", cell_bold), Paragraph("<b>Declared Value on Package</b>", cell_bold), Paragraph("<b>Status</b>", cell_bold)],
        [Paragraph("Rule 6(1)(a) Manufacturer", cell_normal), Paragraph(f"{decl.get('manufacturer_name', 'N/A')}<br/>{decl.get('manufacturer_address', 'N/A')}", cell_normal), Paragraph("Verified" if decl.get('manufacturer_name') else "MISSING", cell_success if decl.get('manufacturer_name') else cell_danger)],
        [Paragraph("Rule 6(1)(b) Commodity Name", cell_normal), Paragraph(str(decl.get('commodity_name', 'N/A')), cell_normal), Paragraph("Verified" if decl.get('commodity_name') else "MISSING", cell_success if decl.get('commodity_name') else cell_danger)],
        [Paragraph("Rule 6(1)(c) Net Quantity", cell_normal), Paragraph(str(decl.get('net_quantity_raw', f"{decl.get('net_quantity_value', '')} {decl.get('net_quantity_unit', '')}")), cell_normal), Paragraph("Verified" if decl.get('net_quantity_raw') or decl.get('net_quantity_value') else "MISSING", cell_success if (decl.get('net_quantity_raw') or decl.get('net_quantity_value')) else cell_danger)],
        [Paragraph("Rule 6(1)(e) Maximum Retail Price (MRP)", cell_normal), Paragraph(str(decl.get('mrp_raw', f"₹ {decl.get('mrp_value', 'N/A')}")), cell_normal), Paragraph("Verified" if decl.get('mrp_raw') or decl.get('mrp_value') else "MISSING", cell_success if (decl.get('mrp_raw') or decl.get('mrp_value')) else cell_danger)],
        [Paragraph("Rule 6(1)(da) Unit Sale Price (USP)", cell_normal), Paragraph(str(decl.get('unit_sale_price', 'Not Declared')), cell_normal), Paragraph("Verified" if decl.get('unit_sale_price') else "MISSING", cell_success if decl.get('unit_sale_price') else cell_danger)],
        [Paragraph("Rule 6(1)(d) Mfg / Pkd Date", cell_normal), Paragraph(str(decl.get('manufacturing_date', 'N/A')), cell_normal), Paragraph("Verified" if decl.get('manufacturing_date') else "MISSING", cell_success if decl.get('manufacturing_date') else cell_danger)],
        [Paragraph("Rule 6(1)(f) Consumer Care", cell_normal), Paragraph(f"Phone: {decl.get('consumer_care_phone', 'N/A')}<br/>Email: {decl.get('consumer_care_email', 'N/A')}", cell_normal), Paragraph("Verified" if (decl.get('consumer_care_phone') or decl.get('consumer_care_email')) else "MISSING", cell_success if (decl.get('consumer_care_phone') or decl.get('consumer_care_email')) else cell_danger)],
        [Paragraph("Rule 6(1)(g) Country of Origin", cell_normal), Paragraph(str(decl.get('country_of_origin', 'N/A')), cell_normal), Paragraph("Verified" if decl.get('country_of_origin') else "MISSING", cell_success if decl.get('country_of_origin') else cell_danger)],
        [Paragraph("Rule 6(1)(h) Batch / Lot No", cell_normal), Paragraph(str(decl.get('batch_number', 'N/A')), cell_normal), Paragraph("Verified" if decl.get('batch_number') else "MISSING", cell_success if decl.get('batch_number') else cell_danger)],
    ]
    
    decl_table = Table(decl_rows, colWidths=[140, 300, 80])
    decl_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#DDE1FF")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(decl_table)
    story.append(Spacer(1, 10))
    
    # 4. Table-II Font Size & Readability Analysis
    story.append(Paragraph("2. Rule 7 & Table-II Font Size & Principal Display Panel (PDP) Analysis", section_heading))
    font_data = scan_data.get("table_ii_font_check", {})
    font_rows = [
        [Paragraph("<b>Parameter</b>", cell_bold), Paragraph("<b>PDP Area</b>", cell_bold), Paragraph("<b>Required Minimum</b>", cell_bold), Paragraph("<b>Measured Value</b>", cell_bold), Paragraph("<b>Result</b>", cell_bold)],
        [Paragraph("Numeral Height (Net Qty / MRP)", cell_normal), Paragraph(f"{font_data.get('panel_area_sq_cm', 140)} cm²", cell_normal), Paragraph(f"{font_data.get('required_numeral_height_mm', 2.0)} mm", cell_normal), Paragraph(f"{font_data.get('measured_numeral_height_mm', 2.4)} mm", cell_normal), Paragraph("PASS" if font_data.get('numeral_pass', True) else "FAIL", cell_success if font_data.get('numeral_pass', True) else cell_danger)],
        [Paragraph("Letter Height (General Text)", cell_normal), Paragraph(f"{font_data.get('panel_area_sq_cm', 140)} cm²", cell_normal), Paragraph(f"{font_data.get('required_letter_height_mm', 1.0)} mm", cell_normal), Paragraph(f"{font_data.get('measured_letter_height_mm', 1.6)} mm", cell_normal), Paragraph("PASS" if font_data.get('letter_pass', True) else "FAIL", cell_success if font_data.get('letter_pass', True) else cell_danger)],
    ]
    font_table = Table(font_rows, colWidths=[150, 90, 100, 100, 80])
    font_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#DDE1FF")),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(font_table)
    story.append(Spacer(1, 10))
    
    # 5. Non-Compliance Violations & Legal Penalty Actions
    violations = scan_data.get("violations", [])
    story.append(Paragraph(f"3. Statutory Violations & Legal Notices Identified ({len(violations)})", section_heading))
    if not violations:
        story.append(Paragraph("<font color='#28A745'><b>✓ No statutory violations detected. Package complies with Legal Metrology (Packaged Commodities) Rules, 2011.</b></font>", cell_normal))
    else:
        v_rows = [[Paragraph("<b>Severity</b>", cell_bold), Paragraph("<b>Rule & Title</b>", cell_bold), Paragraph("<b>Statutory Description & Legal Action</b>", cell_bold)]]
        for v in violations:
            v_sev = v.get("severity", "Major")
            v_rows.append([
                Paragraph(f"<font color='{'#DC3545' if v_sev == 'Critical' else '#E88A1E'}'><b>{v_sev}</b></font>", cell_bold),
                Paragraph(f"<b>{v.get('section', 'Rule 6')}</b><br/>{v.get('title', 'Violation')}", cell_bold),
                Paragraph(f"{v.get('description', '')}<br/><b>Statutory Remedy:</b> {v.get('recommendation', '')}<br/><b>Penalty:</b> {v.get('penalty_clause', '')}", cell_normal)
            ])
        v_table = Table(v_rows, colWidths=[70, 150, 300])
        v_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#FFE8E8")),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#DC3545")),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(v_table)
        
    story.append(Spacer(1, 12))
    # 6. Official Stamp & Certification Block
    cert_block = [
        [Paragraph("<b>OFFICIAL SEAL & SIGNATURE:</b><br/>This digital report is generated by the Automated Legal Metrology Compliance Inspection System under authority of Legal Metrology Act, 2009.", cell_normal),
         Paragraph("<b>INSPECTOR SIGNATURE:</b><br/><br/>____________________________<br/>Authorised Legal Metrology Officer", cell_normal)]
    ]
    cert_table = Table(cert_block, colWidths=[320, 200])
    cert_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F0F4FF")),
        ('BOX', (0, 0), (-1, -1), 1, primary_color),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(KeepTogether(cert_table))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_docx_report(scan_data: Dict[str, Any]) -> io.BytesIO:
    """Generates official Word DOCX inspection report."""
    doc = Document()
    
    title = doc.add_heading("GOVERNMENT OF INDIA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph("Department of Consumer Affairs • Legal Metrology Enforcement Wing\nStatutory Packaged Commodities Compliance Inspection Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("Case Summary", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Inspection ID: {scan_data.get('id', 'N/A')}\n")
    p.add_run(f"Brand: {scan_data.get('brand_name', 'N/A')}\n")
    p.add_run(f"Commodity: {scan_data.get('commodity_name', 'N/A')}\n")
    p.add_run(f"Status: {scan_data.get('compliance_status', 'N/A')}\n")
    p.add_run(f"Compliance Score: {scan_data.get('compliance_score', 0)} / 100\n")
    p.add_run(f"Inspector: {scan_data.get('inspector_name', 'Field Inspector')}\n")
    
    doc.add_heading("Rule 6 Mandatory Declarations", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Rule Section"
    hdr_cells[1].text = "Declared Value"
    hdr_cells[2].text = "Status"
    
    decl = scan_data.get("declarations", {})
    for field, val in decl.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(field).replace("_", " ").title()
        row_cells[1].text = str(val) if val is not None else "N/A"
        row_cells[2].text = "Declared" if val else "Missing"
        
    doc.add_heading("Violations & Legal Penalties", level=1)
    violations = scan_data.get("violations", [])
    if not violations:
        doc.add_paragraph("No statutory violations detected. Fully compliant.")
    else:
        for v in violations:
            vp = doc.add_paragraph()
            vp.add_run(f"[{v.get('severity', 'Major')}] {v.get('title', 'Violation')}\n").bold = True
            vp.add_run(f"Section: {v.get('section', 'Rule 6')}\n")
            vp.add_run(f"Description: {v.get('description', '')}\n")
            vp.add_run(f"Remedy: {v.get('recommendation', '')}\n")
            vp.add_run(f"Penalty Clause: {v.get('penalty_clause', '')}\n")
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_scans_to_csv(scans: List[Dict[str, Any]]) -> io.StringIO:
    """Exports multiple inspection cases to CSV."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "Inspection ID", "Brand Name", "Commodity", "Barcode", "Status", 
        "Compliance Score", "Violations Count", "Critical Violations", 
        "Major Violations", "Inspector", "Created At"
    ])
    for s in scans:
        writer.writerow([
            s.get("id"),
            s.get("brand_name"),
            s.get("commodity_name"),
            s.get("barcode_gtin"),
            s.get("compliance_status"),
            s.get("compliance_score"),
            s.get("violations_count", 0),
            s.get("critical_violations", 0),
            s.get("major_violations", 0),
            s.get("inspector_name"),
            s.get("created_at")
        ])
    output.seek(0)
    return output